"""
Simple DOCX extractor to dump text from a .docx file into a plaintext file.
Usage: python tools/extract_docx.py "path/to/Mô tả BTL2-2.docx" output.txt

Requires: python-docx (pip install python-docx)
"""
import sys
from pathlib import Path

def main():
    if len(sys.argv) < 3:
        print("Usage: python extract_docx.py input.docx output.txt")
        return
    input_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])
    try:
        from docx import Document
    except Exception as e:
        print("Missing dependency 'python-docx'. Install with: pip install python-docx")
        return
    if not input_path.exists():
        print(f"Input not found: {input_path}")
        return
    doc = Document(str(input_path))
    texts = []
    for para in doc.paragraphs:
        texts.append(para.text)
    # also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = '\t'.join(cell.text for cell in row.cells)
            texts.append(row_text)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text('\n'.join(texts), encoding='utf-8')
    print(f"Extracted {len(texts)} paragraphs/rows to {out_path}")

if __name__ == '__main__':
    main()
